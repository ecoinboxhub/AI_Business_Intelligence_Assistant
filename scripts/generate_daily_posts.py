import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as REL_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

POSTS_DIR = "Posts Generator"
SCREENSHOTS_DIR = "submission/screenshots"

os.makedirs(POSTS_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# POST CONTENT DEFINITIONS
# ---------------------------------------------------------------------------

LINKEDIN_POSTS = [
    {
        "title": "Day 1: Eliminating AI Mathematical Hallucinations in Business Intelligence",
        "content": (
            "Most executives are hesitant to rely on conversational AI for financial reporting because standard LLMs "
            "frequently hallucinate numbers. When managing millions in revenue, guessing is not an option.\n\n"
            "NexaSphere solves this with a hybrid architecture. All calculations are handled programmatically by Python's "
            "Pandas engine to guarantee 100 percent mathematical accuracy. The AI model is restricted to interpreting causes, "
            "structuring verified facts, and providing strategic recommendations.\n\n"
            "This screenshot demonstrates our executive dashboard view. Decision-makers instantly see total revenue, profit "
            "margins, return rates, and delivery delays without worrying about fabricated metrics."
        ),
        "image": "01_executive_dashboard.png",
        "caption": "Screenshot Feature Visibility: High-level KPI summary grid displaying verified revenue and margin calculations."
    },
    {
        "title": "Day 2: Instant Regional Revenue & Profit Breakdown",
        "content": (
            "Evaluating performance across multiple retail stores and regions typically requires compiling data across "
            "scattered spreadsheets. NexaSphere simplifies this into a single natural language query.\n\n"
            "By asking 'Which regions generate the highest revenue and profit?', the system instantly executes aggregated "
            "queries across sales records, projects dynamic Recharts visualizations, and highlights underlying profitability drivers.\n\n"
            "Notice how the response separates observed facts from recommendations. Leaders get clear data points alongside actionable guidance."
        ),
        "image": "02_regional_revenue_chart.png",
        "caption": "Screenshot Feature Visibility: Dynamic Grouped Bar Chart comparing regional revenue and profit performance."
    },
    {
        "title": "Day 3: Catching Product Return Problems Before They Grow",
        "content": (
            "Product returns are one of the most underestimated threats to retail profitability. A small group of items "
            "can silently consume margins through reverse logistics, restocking costs, and lost customer trust.\n\n"
            "NexaSphere addresses this by computing exact return rates for every product directly from sales and returns "
            "records using Pandas. Unusually high rates are flagged automatically, and the AI narrative explains likely "
            "drivers such as quality issues, misleading descriptions, or sizing problems.\n\n"
            "The screenshot below shows the return rate analysis in action. Problem products surface immediately instead "
            "of hiding inside quarterly reports."
        ),
        "image": "03_return_rates_analysis.png",
        "caption": "Screenshot Feature Visibility: Return rate ranking table highlighting outlier products with computed percentages."
    },
    {
        "title": "Day 4: Measuring True Marketing ROI Without Spreadsheet Chaos",
        "content": (
            "Marketing teams often struggle to prove impact because campaign data lives in one system and revenue lives "
            "in another. Bringing them together usually means fragile spreadsheets and disputed numbers.\n\n"
            "NexaSphere joins campaign spend with attributed sales and computes return on investment deterministically "
            "before any narrative is written. The AI then explains which campaigns over or under performed and why, using "
            "only verified figures.\n\n"
            "In this screenshot, the donut chart breaks down contribution by campaign. Budget conversations become fact "
            "based rather than opinion based."
        ),
        "image": "04_marketing_roi_donut_chart.png",
        "caption": "Screenshot Feature Visibility: Marketing ROI donut chart showing verified spend versus attributed revenue per campaign."
    },
    {
        "title": "Day 5: Fair Performance Reviews Backed by Verified Numbers",
        "content": (
            "Traditional performance reviews reward whoever sells the most. That approach ignores discounting, returns, "
            "and cost of goods, which means the loudest numbers are not always the healthiest ones.\n\n"
            "NexaSphere evaluates employees on both generated revenue and resulting profit margin, calculated "
            "programmatically from transaction data. Managers get a ranked, drillable view that separates genuine value "
            "creators from high volume low margin sellers.\n\n"
            "The screenshot below captures the performance drilldown. Every figure shown was computed by code, so "
            "coaching conversations start from facts."
        ),
        "image": "05_performance_drilldown.png",
        "caption": "Screenshot Feature Visibility: Performance drilldown ranking employees by verified revenue and profit contribution."
    },
    {
        "title": "Day 6: Your Business Command Center, Now in Your Pocket",
        "content": (
            "Critical decisions rarely wait for the office. By the time a leader returns to their desk, a stockout has "
            "spread, a campaign has overspent, or a delivery issue has escalated.\n\n"
            "The NexaSphere mobile app closes that gap. Built with React Native and Expo, it mirrors the web dashboard "
            "with the same deterministic Pandas calculations behind every figure. Revenue, profit margin, return rates, "
            "inventory health and delivery performance appear as live KPI cards and native charts.\n\n"
            "In the screenshot below, the mobile dashboard presents the complete health of the business in one glance. "
            "Same verified math, now truly portable."
        ),
        "image": "06_mobile_dashboard.png",
        "caption": "Screenshot Feature Visibility: Mobile KPI dashboard rendering verified revenue, margin and operations metrics."
    },
    {
        "title": "Day 7: Asking Your Data a Question Should Feel This Easy",
        "content": (
            "Traditional BI tools demand training: query builders, filter logic, chart configuration. Most managers "
            "simply give up and ask someone else.\n\n"
            "NexaSphere flips the model. The mobile assistant offers guided question chips covering the nine core "
            "management questions, from regional profitability to marketing ROI. Tapping a suggestion runs the full "
            "pipeline: deterministic computation first, then a streamed narrative explaining causes and recommendations "
            "based strictly on those results.\n\n"
            "The screenshot below shows the assistant ready to go. No training required, no formulas, just plain "
            "language in and verified insight out."
        ),
        "image": "07_mobile_assistant_chips.png",
        "caption": "Screenshot Feature Visibility: Guided question chips launching instant analysis in the mobile assistant."
    },
    {
        "title": "Day 8: Series Wrap-Up - Compute First, Narrate Second",
        "content": (
            "Over the past eight days we showed how NexaSphere answers real management questions: where revenue comes "
            "from, which products come back, which campaigns pay off, and who truly drives profit.\n\n"
            "One principle made every answer trustworthy. Calculations are performed programmatically by Pandas before "
            "any text is generated, so the AI narrates facts instead of inventing them. The automated test suite in this "
            "final screenshot keeps that guarantee visible: every analytical path is validated continuously.\n\n"
            "Thank you for following the series. If your organization is tired of arguing about whose spreadsheet is "
            "right, this is what confidence in reporting looks like."
        ),
        "image": "08_pytest_green.png",
        "caption": "Screenshot Feature Visibility: Green test suite confirming continuous validation of all analytical computations."
    }
]

TIKTOK_POSTS = [
    {
        "title": "Day 1: Zero Math Hallucinations",
        "content": "Stop letting AI guess your business numbers. NexaSphere computes exact metrics with zero hallucinations.",
        "image": "01_executive_dashboard.png",
        "caption": "Visibility: Executive dashboard metrics view."
    },
    {
        "title": "Day 2: Instant Business Analytics",
        "content": "Ask questions in plain text and get instant charts with accurate profit numbers across all your stores.",
        "image": "02_regional_revenue_chart.png",
        "caption": "Visibility: Multi-region comparative chart."
    },
    {
        "title": "Day 3: Return Rate Radar",
        "content": "See which products get returned the most in seconds. NexaSphere finds problem items before they drain profits.",
        "image": "03_return_rates_analysis.png",
        "caption": "Visibility: Product return rate ranking view."
    },
    {
        "title": "Day 4: Campaign ROI Revealed",
        "content": "Find out which marketing campaigns actually make money. Exact ROI numbers, no guesswork.",
        "image": "04_marketing_roi_donut_chart.png",
        "caption": "Visibility: Campaign ROI donut chart."
    },
    {
        "title": "Day 5: True Team Rankings",
        "content": "Rank your team by real profit, not just sales. NexaSphere shows who truly drives your bottom line.",
        "image": "05_performance_drilldown.png",
        "caption": "Visibility: Employee performance drilldown."
    },
    {
        "title": "Day 6: BI in Your Pocket",
        "content": "Run your whole business from your phone. NexaSphere puts live revenue KPIs right in your pocket.",
        "image": "06_mobile_dashboard.png",
        "caption": "Visibility: Mobile KPI dashboard view."
    },
    {
        "title": "Day 7: One Tap Answers",
        "content": "Tap one question chip and get a full profit answer instantly. Business intelligence made this simple.",
        "image": "07_mobile_assistant_chips.png",
        "caption": "Visibility: Guided question chips."
    },
    {
        "title": "Day 8: Series Finale",
        "content": "Eight days, one promise: every number verified by code before AI speaks. That is how BI should work.",
        "image": "08_pytest_green.png",
        "caption": "Visibility: Passing test suite."
    }
]

INSTAGRAM_POSTS = [
    {
        "title": "Day 1: Clear Data for Smarter Decisions",
        "content": (
            "Transform how you view business metrics. NexaSphere uses Python to ensure all calculations "
            "are completely accurate before generating executive insights. No math errors, just clarity."
        ),
        "image": "01_executive_dashboard.png",
        "caption": "Visibility: Real-time executive KPI metrics dashboard."
    },
    {
        "title": "Day 2: Spot Profit Opportunities Instantly",
        "content": (
            "Easily compare regional performance and profitability. Ask questions in simple English and receive "
            "verified data paired with clear visual charts in seconds."
        ),
        "image": "02_regional_revenue_chart.png",
        "caption": "Visibility: Interactive regional revenue breakdown."
    },
    {
        "title": "Day 3: Stop Silent Margin Loss",
        "content": (
            "Returned products quietly eat into your margins. NexaSphere scans every sale and highlights items "
            "with unusual return rates, so you can fix quality or sizing issues early. Every figure is computed "
            "with Python first, so the story you read is always backed by exact math."
        ),
        "image": "03_return_rates_analysis.png",
        "caption": "Visibility: Return rate analysis for problem products."
    },
    {
        "title": "Day 4: Know Where Your Budget Goes",
        "content": (
            "Stop wondering where your marketing budget goes. NexaSphere compares campaign spend against real "
            "revenue and ranks your best performers instantly. Ask a question in plain English, get an accurate "
            "chart in seconds. No spreadsheets, no errors, just answers."
        ),
        "image": "04_marketing_roi_donut_chart.png",
        "caption": "Visibility: Marketing ROI breakdown by campaign."
    },
    {
        "title": "Day 5: Fair Reviews Start With Facts",
        "content": (
            "Sales alone never tell the full story. NexaSphere measures every team member on both revenue and "
            "profit, giving managers a fair, verified picture of true performance. One question, one clear "
            "answer, zero manual number crunching."
        ),
        "image": "05_performance_drilldown.png",
        "caption": "Visibility: Employee revenue and profit drilldown."
    },
    {
        "title": "Day 6: Decisions Do Not Wait for the Desk",
        "content": (
            "Your business does not stop when you leave your desk. The NexaSphere mobile app brings the same "
            "verified KPIs to your phone: revenue, margins, return rates and delivery performance, updated live. "
            "Accurate numbers, wherever decisions happen."
        ),
        "image": "06_mobile_dashboard.png",
        "caption": "Visibility: Live mobile KPI dashboard."
    },
    {
        "title": "Day 7: Simple English In, Verified Insight Out",
        "content": (
            "No dashboards to learn, no formulas to write. Open NexaSphere, tap a suggested question like which "
            "campaigns pay off, and read a clear, chart-backed answer in seconds. Simple English in, verified "
            "insight out."
        ),
        "image": "07_mobile_assistant_chips.png",
        "caption": "Visibility: Guided questions in the mobile assistant."
    },
    {
        "title": "Day 8: Thank You for Following the Series",
        "content": (
            "That completes our eight day series. One idea carried it from start to finish: AI should explain "
            "decisions, never invent them. Every metric in NexaSphere is computed deterministically by Python and "
            "protected by an automated test suite. Trust your numbers again."
        ),
        "image": "08_pytest_green.png",
        "caption": "Visibility: Automated tests guarding calculation accuracy."
    }
]

TWITTER_POSTS = [
    {
        "title": "Day 1: Zero Math Error AI",
        "content": "NexaSphere combines exact Python math with AI insights. No numerical hallucinations for your executive reports.",
        "image": "01_executive_dashboard.png",
        "caption": "Visibility: KPI summary cards."
    },
    {
        "title": "Day 2: Fast Business Insights",
        "content": "Ask any business query in natural language and receive exact metrics alongside dynamic visual charts instantly.",
        "image": "02_regional_revenue_chart.png",
        "caption": "Visibility: Regional profit bar chart."
    },
    {
        "title": "Day 3: Return Rate Alerts",
        "content": "High return rates hide in plain sight. NexaSphere scans every product and flags the outliers with exact math.",
        "image": "03_return_rates_analysis.png",
        "caption": "Visibility: Outlier product return table."
    },
    {
        "title": "Day 4: Real Marketing ROI",
        "content": "Which campaigns pay off? NexaSphere computes true marketing ROI from your raw data. No estimates, no hallucinations.",
        "image": "04_marketing_roi_donut_chart.png",
        "caption": "Visibility: Campaign ROI donut chart."
    },
    {
        "title": "Day 5: Profit Over Volume",
        "content": "Top sellers are not always top earners. NexaSphere ranks your team on profit too, with numbers you can verify.",
        "image": "05_performance_drilldown.png",
        "caption": "Visibility: Revenue and profit drilldown."
    },
    {
        "title": "Day 6: KPIs That Travel",
        "content": "Revenue, margins and delivery KPIs now travel with you. NexaSphere mobile puts verified numbers in your pocket.",
        "image": "06_mobile_dashboard.png",
        "caption": "Visibility: Mobile KPI cards."
    },
    {
        "title": "Day 7: BI Like Texting",
        "content": "Tap a suggested question, get a chart-backed answer in seconds. NexaSphere makes BI feel like texting your data.",
        "image": "07_mobile_assistant_chips.png",
        "caption": "Visibility: Question chips."
    },
    {
        "title": "Day 8: Series Wrap-Up",
        "content": "End of our eight day series. One rule held throughout: compute first, narrate second. Trust every number.",
        "image": "08_pytest_green.png",
        "caption": "Visibility: Green test suite."
    }
]

CHARACTER_LIMITS = {
    "LinkedIn": 1500,
    "TikTok": 120,
    "Instagram": 300,
    "Twitter": 130,
}

DEPLOYED_LINKS = {
    "Live Web Dashboard (GitHub Pages)": "https://ecoinboxhub.github.io/AI_Business_Intelligence_Assistant/",
    "Mobile App Release v1.0.0 (Expo bundle)": "https://github.com/ecoinboxhub/AI_Business_Intelligence_Assistant/releases/download/v1.0.0/mobile-release.zip",
    "Source Repository (GitHub)": "https://github.com/ecoinboxhub/AI_Business_Intelligence_Assistant",
}

# ---------------------------------------------------------------------------
# COVER PAGE BUILDER
# ---------------------------------------------------------------------------

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, REL_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_cover_page(doc, platform_name, posts_data):
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("NexaSphere AI Business Intelligence Assistant")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title_para.paragraph_format.space_after = Pt(2)

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(f"Comprehensive Daily {platform_name} Content Suite")
    sub_run.font.size = Pt(14)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(79, 70, 229)
    sub_para.paragraph_format.space_after = Pt(2)

    meta_para = doc.add_paragraph()
    limit = CHARACTER_LIMITS[platform_name]
    meta_run = meta_para.add_run(
        f"{len(posts_data)} posts | Maximum {limit} characters per post | "
        f"Zero emojis | High-resolution screenshot attachments included"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 116, 139)
    meta_para.paragraph_format.space_after = Pt(16)

    about_head = doc.add_paragraph()
    about_head_run = about_head.add_run("About This Implementation")
    about_head_run.font.size = Pt(13)
    about_head_run.font.bold = True
    about_head_run.font.color.rgb = RGBColor(30, 41, 59)
    about_head.paragraph_format.space_before = Pt(6)
    about_head.paragraph_format.space_after = Pt(4)

    intro_1 = doc.add_paragraph()
    intro_1_run = intro_1.add_run(
        "NexaSphere is a local-first, full-stack conversational business intelligence assistant. "
        "Decision-makers ask any of nine core management questions in plain natural language and receive "
        "structured answers combining dynamic charts, verified findings and strategic recommendations."
    )
    intro_1_run.font.size = Pt(10.5)
    intro_1_run.font.color.rgb = RGBColor(51, 65, 85)
    intro_1.paragraph_format.space_after = Pt(8)

    intro_2 = doc.add_paragraph()
    intro_2_run = intro_2.add_run(
        "Every number in every answer is computed deterministically by Python's Pandas analytics engine "
        "before the AI model writes a single word of narrative. The language model explains causes and "
        "recommendations but never performs math, eliminating numerical hallucinations entirely."
    )
    intro_2_run.font.size = Pt(10.5)
    intro_2_run.font.color.rgb = RGBColor(51, 65, 85)
    intro_2.paragraph_format.space_after = Pt(10)

    highlights_head = doc.add_paragraph()
    highlights_head_run = highlights_head.add_run("Implementation Highlights")
    highlights_head_run.font.size = Pt(13)
    highlights_head_run.font.bold = True
    highlights_head_run.font.color.rgb = RGBColor(30, 41, 59)
    highlights_head.paragraph_format.space_after = Pt(4)

    highlights = [
        "Backend API: FastAPI with a DuckDB analytical engine, risk scoring, what-if simulation and SSE streamed responses on port 5050.",
        "Web Dashboard: React 18 (Vite) executive interface with Recharts visualizations across Dashboard, Assistant and Performance pages.",
        "Mobile App: Expo / React Native companion app mirroring live KPI cards and the guided-question assistant on iOS and Android.",
        "Data flexibility: sales, campaigns, deliveries, inventory and targets load from CSV, TSV, Excel, Parquet or JSON without code changes.",
        "Quality assurance: a 60-test automated suite validates analytical accuracy benchmarks and data-source resolution on every commit.",
    ]
    for line in highlights:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet_run = bullet.add_run(line)
        bullet_run.font.size = Pt(10.5)
        bullet_run.font.color.rgb = RGBColor(51, 65, 85)
        bullet.paragraph_format.space_after = Pt(4)

    links_head = doc.add_paragraph()
    links_head_run = links_head.add_run("Deployed and Published Links")
    links_head_run.font.size = Pt(13)
    links_head_run.font.bold = True
    links_head_run.font.color.rgb = RGBColor(30, 41, 59)
    links_head.paragraph_format.space_before = Pt(10)
    links_head.paragraph_format.space_after = Pt(4)

    for label, url in DEPLOYED_LINKS.items():
        link_para = doc.add_paragraph(style="List Bullet")
        label_run = link_para.add_run(f"{label}: ")
        label_run.font.size = Pt(10.5)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(51, 65, 85)
        add_hyperlink(link_para, url, url)
        link_para.paragraph_format.space_after = Pt(4)

    ci_note = doc.add_paragraph()
    ci_note_run = ci_note.add_run(
        "Continuous delivery: GitHub Actions runs the full test suite, publishes the web dashboard to GitHub Pages "
        "on every push to main, and packages the mobile release automatically on version tags."
    )
    ci_note_run.font.size = Pt(10)
    ci_note_run.font.italic = True
    ci_note_run.font.color.rgb = RGBColor(100, 116, 139)
    ci_note.paragraph_format.space_before = Pt(8)
    ci_note.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

# ---------------------------------------------------------------------------
# DOCX BUILDER FUNCTION
# ---------------------------------------------------------------------------
# CHARACTER LIMIT VALIDATION (strict platform conformance)
# ---------------------------------------------------------------------------

def validate_post_lengths(platform_name, posts_data):
    limit = CHARACTER_LIMITS[platform_name]
    for item in posts_data:
        length = len(item["content"])
        if length > limit:
            raise ValueError(
                f"{platform_name} post '{item['title']}' is {length} characters "
                f"(limit {limit}). Shorten the content before generating."
            )

# ---------------------------------------------------------------------------
# DOCX BUILDER FUNCTION
# ---------------------------------------------------------------------------

def create_platform_docx(filename, platform_name, posts_data):
    validate_post_lengths(platform_name, posts_data)

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Cover page introducing the full implementation
    build_cover_page(doc, platform_name, posts_data)

    for item in posts_data:
        # Post Title
        t_para = doc.add_paragraph()
        t_run = t_para.add_run(item["title"])
        t_run.font.size = Pt(13)
        t_run.font.bold = True
        t_run.font.color.rgb = RGBColor(30, 41, 59)
        t_para.paragraph_format.space_before = Pt(10)
        t_para.paragraph_format.space_after = Pt(4)

        # Post Content
        c_para = doc.add_paragraph()
        c_run = c_para.add_run(item["content"])
        c_run.font.size = Pt(10.5)
        c_run.font.color.rgb = RGBColor(51, 65, 85)
        c_para.paragraph_format.space_after = Pt(8)

        # Embedded Image (if available)
        image_path = os.path.join(SCREENSHOTS_DIR, item["image"])
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(5.5))

            cap_para = doc.add_paragraph()
            cap_run = cap_para.add_run(item["caption"])
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            cap_run.font.color.rgb = RGBColor(100, 116, 139)
            cap_para.paragraph_format.space_after = Pt(16)
        else:
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(f"[Screenshot Attachment: {item['image']}]")
            note_run.font.size = Pt(9)
            note_run.font.italic = True

    filepath = os.path.join(POSTS_DIR, filename)
    doc.save(filepath)
    print(f"Successfully generated: {filepath}")

if __name__ == "__main__":
    create_platform_docx("LinkedIn_Posts.docx", "LinkedIn", LINKEDIN_POSTS)
    create_platform_docx("TikTok_Posts.docx", "TikTok", TIKTOK_POSTS)
    create_platform_docx("Instagram_Posts.docx", "Instagram", INSTAGRAM_POSTS)
    create_platform_docx("Twitter_Posts.docx", "Twitter", TWITTER_POSTS)
